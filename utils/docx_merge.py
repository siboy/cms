"""
Merge chunk-chunk (HTML yang sudah diedit user) kembali menjadi .docx.

Strategi: bangun dokumen baru dari template kosong (atau template original
docx sebagai style-source), lalu inject chunk sesuai urutan order_idx.
Parser HTML sederhana via BeautifulSoup: h1..h6, p, ul/ol/li, strong, em,
u, br, img, table/tr/td didukung. Atribut lain diabaikan.
"""
from __future__ import annotations

import os
import shutil
from copy import deepcopy
from typing import Optional

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches


HEADING_TAGS = {f"h{i}": i for i in range(1, 7)}


def _add_runs_from_node(paragraph, node, media_map: dict[str, str]):
    """Recurse ke children node → tambah runs ke paragraph."""
    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return

    tag = node.name
    if tag == "br":
        paragraph.add_run().add_break()
        return

    if tag == "img":
        rid_or_path = node.get("data-rid") or node.get("src") or ""
        img_path = media_map.get(rid_or_path) or rid_or_path
        if img_path and os.path.isfile(img_path):
            try:
                paragraph.add_run().add_picture(img_path, width=Inches(5.5))
            except Exception:
                paragraph.add_run(f"[image:{os.path.basename(img_path)}]")
        return

    # inline formatting
    if tag in ("strong", "b"):
        for c in node.children:
            if isinstance(c, NavigableString):
                r = paragraph.add_run(str(c))
                r.bold = True
            else:
                _add_runs_from_node(paragraph, c, media_map)
        return
    if tag in ("em", "i"):
        for c in node.children:
            if isinstance(c, NavigableString):
                r = paragraph.add_run(str(c))
                r.italic = True
            else:
                _add_runs_from_node(paragraph, c, media_map)
        return
    if tag == "u":
        for c in node.children:
            if isinstance(c, NavigableString):
                r = paragraph.add_run(str(c))
                r.underline = True
            else:
                _add_runs_from_node(paragraph, c, media_map)
        return

    # fallback: recurse
    for c in node.children:
        _add_runs_from_node(paragraph, c, media_map)


def _render_block(doc, node, media_map: dict[str, str]):
    """Render block-level node (h1..h6, p, ul, ol, table) ke document."""
    tag = node.name

    if tag in HEADING_TAGS:
        p = doc.add_paragraph(style=f"Heading {HEADING_TAGS[tag]}")
        _add_runs_from_node(p, node, media_map)
        return

    if tag == "p":
        p = doc.add_paragraph()
        _add_runs_from_node(p, node, media_map)
        return

    if tag in ("ul", "ol"):
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
            _add_runs_from_node(p, li, media_map)
        return

    if tag == "li":
        p = doc.add_paragraph(style="List Bullet")
        _add_runs_from_node(p, node, media_map)
        return

    if tag == "table":
        rows = node.find_all("tr")
        if not rows:
            return
        n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"
        for ri, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            for ci, td in enumerate(cells):
                cell = table.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_runs_from_node(p, td, media_map)
        return

    # fallback: wrap in paragraph
    p = doc.add_paragraph()
    _add_runs_from_node(p, node, media_map)


def merge_chunks_to_docx(
    chunks: list[dict],
    out_path: str,
    template_docx: Optional[str] = None,
    media_map: Optional[dict[str, str]] = None,
):
    """
    chunks: list of dict, each {order_idx, heading_level, heading_text,
            content_html, ...} sorted by order_idx.
    template_docx: path ke docx asli — dipakai sebagai template (style source).
                   Kalau None, Document() default.
    media_map: {rid_atau_src: absolute_path_file_gambar}

    Return out_path.
    """
    media_map = media_map or {}

    if template_docx and os.path.isfile(template_docx):
        tmp_template = out_path + ".template.docx"
        shutil.copy(template_docx, tmp_template)
        doc = Document(tmp_template)
        # Clear body existing content, keep styles
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)
    else:
        doc = Document()

    sorted_chunks = sorted(chunks, key=lambda c: c["order_idx"])
    for ch in sorted_chunks:
        html = ch.get("content_html") or ""
        if not html.strip():
            continue
        soup = BeautifulSoup(html, "html.parser")
        for child in soup.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    p = doc.add_paragraph()
                    p.add_run(txt)
            elif child.name:
                _render_block(doc, child, media_map)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    tmp_template = out_path + ".template.docx"
    if os.path.exists(tmp_template):
        os.remove(tmp_template)
    return out_path
